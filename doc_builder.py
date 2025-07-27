# doc_builder.py

from docx import Document
import json
from io import BytesIO

def build_word_doc(summaries, model_purpose, inputs_data, outputs_data, logic_steps, checks_data, assumptions_text):
    doc = Document()
    doc.add_heading("Named Range JSON Summary", 0)

    for name, summary in summaries.items():
        doc.add_heading(name, level=1)
        for key, value in summary.items():
            if isinstance(value, (list, dict)):
                value = json.dumps(value, indent=2)
            doc.add_paragraph(f"{key}: {value}")

    doc.add_page_break()
    doc.add_heading("📄 Spreadsheet Documentation", 0)

    # Version Control
    doc.add_heading("Version Control", level=1)
    doc.add_heading("Model Version Control", level=2)
    for col in ["Version", "Date", "Info", "Updated by", "Reviewed by", "Review Date"]:
        doc.add_paragraph(f"{col}: __________")

    doc.add_heading("Documentation Version Control", level=2)
    for col in ["Version", "Date", "Info", "Updated by", "Reviewed by", "Review Date"]:
        doc.add_paragraph(f"{col}: __________")

    # Ownership
    doc.add_heading("Ownership", level=1)
    doc.add_paragraph("Owner: __________")
    doc.add_paragraph("Risk rating (or other client control standard): __________")
    doc.add_paragraph("Internal audit history: __________")

    # Purpose
    doc.add_heading("Purpose", level=1)
    doc.add_paragraph(model_purpose)

    # Inputs
    doc.add_heading("Inputs", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[:] = [c.text for c in ["No.", "Name", "Type", "Source", "Info"]]

    for row in inputs_data:
        r = table.add_row().cells
        r[0].text = str(row["No."])
        r[1].text = row["Name"]
        r[2].text = row["Type"]
        r[3].text = row["Source"]
        r[4].text = row["Info"]

    # Outputs
    doc.add_heading("Outputs", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[:] = [c.text for c in ["No.", "Name", "Description"]]

    for row in outputs_data:
        r = table.add_row().cells
        r[0].text = str(row["No."])
        r[1].text = row["Name"]
        r[2].text = row["Description"]

    # Logic
    doc.add_heading("Logic", level=1)
    if logic_steps:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[:] = [c.text for c in ["Step", "Named Range", "Description"]]

        for row in logic_steps:
            r = table.add_row().cells
            r[0].text = str(row["Step"])
            r[1].text = row["Named Range"]
            r[2].text = row["Description"]
    else:
        doc.add_paragraph("⚠ No logic components found using `_cN_` naming pattern.")

    # Checks
    doc.add_heading("Checks and Validation", level=1)
    if checks_data:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[:] = [c.text for c in ["Check No.", "Named Range", "Description"]]

        for row in checks_data:
            r = table.add_row().cells
            r[0].text = str(row["Check No."])
            r[1].text = row["Named Range"]
            r[2].text = row["Description"]
    else:
        doc.add_paragraph("⚠ No validation checks found using `_chN_` naming pattern.")

    # Assumptions
    doc.add_heading("Assumptions and Limitations", level=1)
    doc.add_paragraph(assumptions_text)

    # TAS
    doc.add_heading("TAS Compliance", level=1)
    doc.add_paragraph("Describe how the model complies with TAS:")

    # Return file-like object
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output
