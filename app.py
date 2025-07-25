import streamlit as st
from openpyxl import load_workbook
from openpyxl.utils import column_index_from_string, get_column_letter
from io import BytesIO
import re
import os
from collections import defaultdict
import graphviz
from docx import Document
import pandas as pd

st.set_page_config(page_title="Named Range Formula Remapper", layout="wide")

# ✅ Step 1: Add this here
if "json_summaries" not in st.session_state:
    st.session_state.json_summaries = None

st.title("\U0001F4D8 AI Powered Spreadsheet Documentation")

if "expanded_all" not in st.session_state:
    st.session_state.expanded_all = False

def toggle():
    st.session_state.expanded_all = not st.session_state.expanded_all

st.button("🔁 Expand / Collapse All Named Ranges", on_click=toggle)

# Allow manual mapping of external references like [1], [2], etc.

with st.expander("🔧 Manual Mapping for External References", expanded=False):
    st.subheader("Manual Mapping for External References")
    external_refs = {}
    for i in range(1, 10):
        ref_key = f"[{i}]"
        workbook_name = st.text_input(f"Map external reference {ref_key} to workbook name (e.g., Mortality_Model_Inputs.xlsx)", key=ref_key)
        if workbook_name:
            external_refs[ref_key] = workbook_name

uploaded_files = st.file_uploader("\U0001F4C2 Upload Excel files", type=["xlsx"], accept_multiple_files=True)

if uploaded_files:
    all_named_cell_map = {}
    all_named_ref_info = {}
    file_display_names = {}
    named_ref_formulas = {}

    for uploaded_file in uploaded_files:
        display_name = uploaded_file.name
        file_display_names[display_name] = uploaded_file
        st.header(f"\U0001F4C4 File: {display_name}")
        wb = load_workbook(BytesIO(uploaded_file.read()), data_only=False)

        for name in wb.defined_names:
            dn = wb.defined_names[name]
            if dn.is_external or not dn.attr_text:
                continue
            for sheet_name, ref in dn.destinations:
                try:
                    ws = wb[sheet_name]
                    ref_clean = ref.replace("$", "").split("!")[-1]
                    cells = ws[ref_clean] if ":" in ref_clean else [[ws[ref_clean]]]

                    min_row = min(cell.row for row in cells for cell in row)
                    min_col = min(cell.column for row in cells for cell in row)

                    coord_set = set()
                    for row in cells:
                        for cell in row:
                            r, c = cell.row, cell.column
                            row_offset = r - min_row + 1
                            col_offset = c - min_col + 1
                            all_named_cell_map[(display_name, sheet_name, r, c)] = (name, row_offset, col_offset)
                            coord_set.add((r, c))
                    all_named_ref_info[name] = (display_name, sheet_name, coord_set, min_row, min_col)
                except:
                    continue

    def remap_formula(formula, current_file, current_sheet):
        if not formula:
            return ""

        def cell_address(row, col):
            return f"{get_column_letter(col)}{row}"

        def remap_single_cell(ref, default_file, default_sheet):
        # Use regex to safely extract sheet name and address
            match = re.match(r"(?:'([^']+)'|([^'!]+))!([$A-Z]+[0-9]+)", ref)
            if match:
                sheet_name = match.group(1) or match.group(2)
                addr = match.group(3)
                # Check for external reference like [1]
                external_match = re.match(r"\[(\d+)\]", sheet_name)
                if external_match:
                    external_ref = external_match.group(0)
                    external_file = external_refs.get(external_ref, external_ref)
                    return f"[{external_file}]{ref}"
            else:
                sheet_name = default_sheet
                addr = ref

            addr = addr.replace("$", "").upper()
            match = re.match(r"([A-Z]+)([0-9]+)", addr)
            if not match:
                return ref
            col_str, row_str = match.groups()
            row = int(row_str)
            col = column_index_from_string(col_str)

            key = (default_file, sheet_name, row, col)
            if key in all_named_cell_map:
                name, r_off, c_off = all_named_cell_map[key]
                return f"[{default_file}]{name}[{r_off}][{c_off}]"
            else:
                return f"{sheet_name}!{addr}"  # Removed [default_file] from fallback

        def remap_range(ref, default_file, default_sheet):
            # Handle external references like [1]Sheet!A1
            if ref.startswith("["):
                match = re.match(r"\[(\d+)\]", ref)
                if match:
                    external_ref = match.group(0)
                    external_file = external_refs.get(external_ref, external_ref)
                    return f"[{external_file}]{ref}"

            # Use regex to safely extract sheet name and address or range
            match = re.match(r"(?:'([^']+)'|([^'!]+))!([$A-Z]+[0-9]+(?::[$A-Z]+[0-9]+)?)", ref)
            if match:
                sheet_name = match.group(1) or match.group(2)
                addr = match.group(3)
            else:
                sheet_name = default_sheet
                addr = ref

            addr = addr.replace("$", "").upper()
            if ":" not in addr:
                return remap_single_cell(ref, default_file, default_sheet)

            # It's a range like A1:B2
            start, end = addr.split(":")
            m1 = re.match(r"([A-Z]+)([0-9]+)", start)
            m2 = re.match(r"([A-Z]+)([0-9]+)", end)
            if not m1 or not m2:
                return ref
            start_col = column_index_from_string(m1.group(1))
            start_row = int(m1.group(2))
            end_col = column_index_from_string(m2.group(1))
            end_row = int(m2.group(2))

            label_set = set()
            for row in range(start_row, end_row + 1):
                for col in range(start_col, end_col + 1):
                    key = (default_file, sheet_name, row, col)
                    if key in all_named_cell_map:
                        name, r_off, c_off = all_named_cell_map[key]
                        label_set.add(f"[{default_file}]{name}[{r_off}][{c_off}]")
                    else:
                        label_set.add(f"{sheet_name}!{get_column_letter(col)}{row}")
            return ", ".join(sorted(label_set))

        pattern = r"(?<![A-Za-z0-9_])(?:'[^']+'|[A-Za-z0-9_]+)!\$?[A-Z]{1,3}\$?[0-9]{1,7}(?::\$?[A-Z]{1,3}\$?[0-9]{1,7})?|(?<![A-Za-z0-9_])\$?[A-Z]{1,3}\$?[0-9]{1,7}(?::\$?[A-Z]{1,3}\$?[0-9]{1,7})?"
        matches = list(re.finditer(pattern, formula))
        replaced_formula = formula
        offset = 0
        for match in matches:
            raw = match.group(0)
            remapped = remap_range(raw, current_file, current_sheet)
            start, end = match.start() + offset, match.end() + offset
            replaced_formula = replaced_formula[:start] + remapped + replaced_formula[end:]
            offset += len(remapped) - len(raw)
        return replaced_formula

    for (name, (file_name, sheet_name, coord_set, min_row, min_col)) in all_named_ref_info.items():
        entries = []
        formulas_for_graph = []

        try:
            file_bytes = file_display_names[file_name]
            wb = load_workbook(BytesIO(file_bytes.getvalue()), data_only=False)
            ws = wb[sheet_name]
            min_col_letter = get_column_letter(min([c for (_, c) in coord_set]))
            max_col_letter = get_column_letter(max([c for (_, c) in coord_set]))
            min_row_num = min([r for (r, _) in coord_set])
            max_row_num = max([r for (r, _) in coord_set])
            ref_range = f"{min_col_letter}{min_row_num}:{max_col_letter}{max_row_num}"
            cell_range = ws[ref_range] if ":" in ref_range else [[ws[ref_range]]]

            for row in cell_range:
                for cell in row:
                    row_offset = cell.row - min_row + 1
                    col_offset = cell.column - min_col + 1
                    label = f"{name}[{row_offset}][{col_offset}]"

                    try:
                        formula = None
                        if isinstance(cell.value, str) and cell.value.startswith("="):
                            formula = cell.value.strip()
                        elif hasattr(cell, 'value') and hasattr(cell.value, 'text'):
                            formula = str(cell.value.text).strip()
                        elif hasattr(cell, 'value'):
                            formula = str(cell.value)

                        if formula:
                            remapped = remap_formula(formula, file_name, sheet_name)
                            formulas_for_graph.append(remapped)
                        elif cell.value is not None:
                            formula = f"[value] {str(cell.value)}"
                            remapped = formula
                        else:
                            formula = "(empty)"
                            remapped = formula
                    except Exception as e:
                        formula = f"[error reading cell: {e}]"
                        remapped = formula

                    entries.append(f"{label} = {formula}\n → {remapped}")
        except Exception as e:
            entries.append(f"❌ Error accessing {name} in {sheet_name}: {e}")

        named_ref_formulas[name] = formulas_for_graph
        
        limit = 50
        snippet = entries if limit is None else entries[:limit]
        
        with st.expander(
            f"📌 Named Range: {name} → {sheet_name} in {file_name}",
            expanded=st.session_state.expanded_all
        ):
            st.code("\n".join(snippet), language="text")
            if limit is not None and len(entries) > limit:
                st.write(f"...and {len(entries) - limit} more lines hidden")
                
    # —– Missing direct cell references (not in any named range) —–
    with st.expander("⚠️ Missing Direct Cell References", expanded=True):
        st.markdown("#### 🔍 Check for A1-style cell references not covered by any named range")

        raw_ref_re = re.compile(r"\b([A-Z]{1,3}[0-9]{1,7})\b")
        missing_refs = defaultdict(set)

        for nm, formulas in named_ref_formulas.items():
            for f in formulas:
                for ref in raw_ref_re.findall(f):
                    if re.search(rf"\[{nm}\]\[\d+\]\[\d+\]", f):
                        continue
                    missing_refs[nm].add(ref)

        if missing_refs:
            for nm, refs in missing_refs.items():
                st.warning(
                    f"In **{nm}**, these direct cell refs weren’t wrapped by a named range: "
                    f"{', '.join(sorted(refs))}"
                )
        else:
            st.success("✅ No missing direct cell references found.")
    
    # Dependency Graph
    st.subheader("🔗 Dependency Graph")
    dot = graphviz.Digraph()
    dot.attr(compound='true', rankdir='LR')

    grouped = defaultdict(list)
    for name, (file, *_rest) in all_named_ref_info.items():
        grouped[file].append(name)

    dependencies = defaultdict(set)
    for target, formulas in named_ref_formulas.items():
        joined = " ".join(formulas)
        for source in named_ref_formulas:
            if source != target and re.search(rf"\b{re.escape(source)}\b", joined):
                dependencies[target].add(source)

    for i, (file_name, nodes) in enumerate(grouped.items()):
        with dot.subgraph(name=f"cluster_{i}") as c:
            c.attr(label=file_name)
            c.attr(style='filled', color='lightgrey')
            for node in nodes:
                c.node(node)

    for target, sources in dependencies.items():
        for source in sources:
            dot.edge(source, target)

    st.graphviz_chart(dot)
# --- JSON Summary Generation Section ---
    
    st.subheader("🧠 Generate JSON and Documentation")

    generate_json = st.button("🧾 Generate")

    if generate_json:
        from openai import OpenAI
        import json

        client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

        summaries = {}

        for name, formulas in named_ref_formulas.items():
            if not formulas:
                continue

            prompt = f"""
    You are an expert actuary and spreadsheet analyst.
   
    Given the following remapped formulas from an Excel named range, summarize the pattern behind the calculations in a general form.
    Each formula follows a remapped structure using notation like [1][2] to indicate row and column indices.

    Please return a JSON object like:
    {{
      "file_name": "MyWorkbook.xlsx",
      "sheet_name": "Inputs",
      "excel_range": "B2:D5",
      "named_range": "MyNamedRange",
      "summary": "Description of what the formula does",
      "general_formula": "for i in range(...): for j in range(...): Result[i][j] = ...",
      "dependencies": ["OtherNamedRange1", "OtherNamedRange2"],
      "notes": "Any caveats, limitations, or variations found"
    }}

    Formulas:
    {formulas[:10]}  # Sample first 10 for context

    Only return the JSON.
    """

            try:
                response = client.chat.completions.create(
                    model="gpt-4",
                    messages=[
                        {"role": "system", "content": "You summarize spreadsheet formulas into structured JSON."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.3
                )
                content = response.choices[0].message.content
                parsed = json.loads(content)
                
                # Get file_name, sheet_name, coord_set for this named range
                file_name, sheet_name, coord_set, *_ = all_named_ref_info[name]

                # Calculate Excel range
                min_col_letter = get_column_letter(min([c for (_, c) in coord_set]))
                max_col_letter = get_column_letter(max([c for (_, c) in coord_set]))
                min_row_num = min([r for (r, _) in coord_set])
                max_row_num = max([r for (r, _) in coord_set])
                excel_range = f"{min_col_letter}{min_row_num}:{max_col_letter}{max_row_num}"
                
                parsed["named_range"] = name  
                parsed["file_name"] = file_name
                parsed["sheet_name"] = sheet_name
                parsed["excel_range"] = excel_range
                parsed["dependencies"] = sorted(dependencies.get(name, []))
                summaries[name] = parsed
            except Exception as e:
                summaries[name] = {"named_range": name,"error": str(e)}

        with st.expander("📦 View JSON Output", expanded=False):
            st.json(summaries)

        #prepare content for documentation

        ## -----Hints----###

        hint_keywords = set()

        for name in summaries:
            name_lower = name.lower()
            if "ax" in name_lower:
                hint_keywords.add("annuity rates")
            if "qx" in name_lower or "mortality" in name_lower:
                hint_keywords.add("mortality rates")
            if "sx" in name_lower:
                hint_keywords.add("survival probabilities")
            if "stoch" in name_lower or "rand" in name_lower or "stochastic" in name_lower:
                hint_keywords.add("simulation-based projections")
            if "vol" in name_lower or "sd" in name_lower or "sigma" in name_lower:
                hint_keywords.add("volatility inputs or stochastic variation")
            if "drift" in name_lower:
                hint_keywords.add("long-term mortality trends or drift terms")
            if "kapp" in name_lower or "beta" in name_lower or "alpha" in name_lower:
                hint_keywords.add("Lee-Carter model parameters")

        # Compose final sentence
        if hint_keywords:
            hint_sentence = "This model work with " + ", ".join(sorted(hint_keywords)) + "."
        else:
            hint_sentence = ""
        
        # --- Generate high-level Purpose description ---
        try:
            joined_descriptions = "\n".join(
                f"{k}: {v.get('summary', '')}" for k, v in summaries.items() if "summary" in v
            )
            joined_formulas = "\n".join(
                f"{k}: {v.get('general_formula', '')}" for k, v in summaries.items() if "general_formula" in v
            )

            purpose_prompt = f"""You are an expert actuary and spreadsheet modeller.

        You are reviewing an Excel model based on the **Lee-Carter mortality framework**.

        {hint_sentence}

        The model uses named ranges and formulas structured to perform actuarial calculations.

        Below are descriptions of how various parts of the model behave:

        --- Summaries ---
        {joined_descriptions}

        --- Formula patterns ---
        {joined_formulas}

        Using this information, write a **concise and confident purpose statement** for documentation. Your paragraph should follow this structure:

        1. Start with a clear sentence about what the model is designed to do (e.g. project mortality, simulate survival rates).
        2. Describe what kinds of inputs it uses (e.g. mortality trends, drift terms, random simulations).
        3. Summarize the types of outputs produced (e.g. annuity rates, survival curves).
        4. Close with a sentence explaining what this model is useful for — pricing, forecasting, risk management, etc.

        Use actuarial language. Do not say “likely”, “possibly”, or “may”. Be direct and factual.
        """

            purpose_response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You write purpose sections for actuarial models."},
                    {"role": "user", "content": purpose_prompt}
                ],
                temperature=0.3
            )

            model_purpose = purpose_response.choices[0].message.content.strip()

        except Exception as e:
            model_purpose = f"Error generating purpose: {e}"

        
        # input data
        input_summaries = {k: v for k, v in summaries.items() if k.startswith("i_")}
        inputs_data = []

        for idx, (name, summary) in enumerate(input_summaries.items(), start=1):
            excel_range = summary.get("excel_range", "")

            # Determine source based on name
            if "_a_" in name:
                source = "Assumptions team"
            elif "_m_" in name:
                source = "Modelling team"
            else:
                source = "Unknown"

            # Determine type based on excel_range
            cell_type = "Unknown"
            if ":" not in excel_range:
                cell_type = "Error"  # invalid or incomplete range
            else:
                try:
                    from_cell, to_cell = excel_range.split(":")
                    if from_cell == to_cell:
                        cell_type = "Cell"
                    else:
                        from_col = re.sub(r"\d", "", from_cell)
                        from_row = int(re.sub(r"\D", "", from_cell))
                        to_col = re.sub(r"\d", "", to_cell)
                        to_row = int(re.sub(r"\D", "", to_cell))

                        if from_col == to_col:
                            cell_type = "Vector"  # vertical
                        elif from_row == to_row:
                            cell_type = "Vector"  # horizontal
                        else:
                            cell_type = "Table"
                except Exception:
                    cell_type = "Error"

            inputs_data.append({
                "No.": idx,
                "Name": name,
                "Type": cell_type,
                "Source": source,
                "Info": ""  # To be filled by GPT
            })
            
        # GPT to populate Info field
        for row in inputs_data:
            input_name = row["Name"]
            summary_json = input_summaries[input_name]
            json_summary = summary_json.get("summary", "")
            general_formula = summary_json.get("general_formula", "")
            sheet = summary_json.get("sheet_name", "")
            excel_range = summary_json.get("excel_range", "")
            prompt = f"""You are an expert actuary and survival modeller.

        You are reviewing a spreadsheet model based on the Lee-Carter mortality model or a closely related framework.

        {hint_sentence}

        You're now documenting the spreadsheet input named `{input_name}`, located in sheet `{sheet}`, cell range `{excel_range}`.

        Its name suggests it's related to: "{input_name}"

        Here is the description of how this input is used in the model:
        "{json_summary}"

        And here is the general formula pattern that references it:
        "{general_formula}"

        Based on all the above, write a concise, confident description of what `{input_name}` represents and how it contributes to the model.

        Use actuarial language. Avoid vague expressions like “might”, “somewhat”, “typically”, or filler phrases like “plays a crucial role” or “is important”. Do not describe patterns in the data (e.g., “decreasing linearly”) unless they are explicitly mentioned.

        Respond with one precise sentence, or two if the second adds new technical detail or context.
"""

            try:
                response = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {"role": "system", "content": "You provide concise descriptions of actuarial inputs."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.3
                )
                row["Info"] = response.choices[0].message.content.strip()
            except Exception as e:
                row["Info"] = f"Error: {e}"
        inputs_df = pd.DataFrame(inputs_data)

        # --- Output data ---
        output_summaries = {k: v for k, v in summaries.items() if k.startswith("o_")}
        outputs_data = []
        for idx, (name, summary) in enumerate(output_summaries.items(), start=1):
            outputs_data.append({
                "No.": idx,
                "Name": name,
                "Description": ""  # To be filled by GPT
            })
        output_names = " ".join(output_summaries.keys()).lower()

                   
        # GPT to populate output descriptions
        for row in outputs_data:
            output_name = row["Name"]
            summary_json = output_summaries[output_name]
            json_summary = summary_json.get("summary", "")
            general_formula = summary_json.get("general_formula", "")
            sheet = summary_json.get("sheet_name", "")
            excel_range = summary_json.get("excel_range", "")

            prompt = f"""You are an expert actuary and survival modeller.

        You are reviewing a spreadsheet output named `{output_name}`, located in sheet `{sheet}`, cell range `{excel_range}`.

        Here is the description of how this output is used in the model:
        "{json_summary}"

        And here is the general formula pattern that defines it:
        "{general_formula}"

        This model is based on the **Lee-Carter mortality model** or a related survival framework.

        Describe in clear actuarial language what `{output_name}` represents and its role in the output of the model. Be confident and specific. Avoid words like "might", "likely", or "possibly".

        Respond with 1–2 precise sentences."""

            try:
                response = client.chat.completions.create(
                    model="gpt-4",
                    messages=[
                        {"role": "system", "content": "You describe actuarial spreadsheet outputs."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.3
                )
                row["Description"] = response.choices[0].message.content.strip()
            except Exception as e:
                row["Description"] = f"Error: {e}"

        outputs_df = pd.DataFrame(outputs_data)

        # --- Logic documentation based on _c1_, _c2_, etc. ---
        logic_summaries = {}
        logic_pattern = re.compile(r"^_c(\d+)_.*")  # matches _c1_, _c2_, etc.

        # Extract and sort logic blocks
        for name in summaries:
            match = logic_pattern.match(name)
            if match:
                step_number = int(match.group(1))
                logic_summaries[step_number] = name

        logic_steps = []
        for step_number in sorted(logic_summaries):
            name = logic_summaries[step_number]
            summary_json = summaries[name]
            json_summary = summary_json.get("summary", "")
            general_formula = summary_json.get("general_formula", "")
            dependencies_list = summary_json.get("dependencies", [])
            excel_range = summary_json.get("excel_range", "")
            sheet = summary_json.get("sheet_name", "")

            prompt = f"""You are an expert actuary and spreadsheet modeller.

        You are reviewing a calculation step in an Excel-based actuarial model built on the Lee-Carter mortality framework.

        The named range for this step is `{name}` (step {step_number}), and it represents a key stage in the spreadsheet's logic.

        Here is the general description of this calculation step:
        "{json_summary}"

        And here is the general formula logic:
        "{general_formula}"

        This step depends on the following named ranges:
        {', '.join(dependencies)}

        Please write 2–3 precise and confident sentences that explain:

        1. The purpose of this calculation step in the model.
        2. What is the calculations and how it contributes to the projection.
        3. Any key inputs or dependencies used in this step.

        Avoid vague language like 'might' or 'possibly' — be concise and clear.

        Respond with 1-2 precise sentences."""

            try:
                response = client.chat.completions.create(
                    model="gpt-4",
                    messages=[
                        {"role": "system", "content": "You describe logic steps in actuarial models clearly."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.3
                )
                explanation = response.choices[0].message.content.strip()
            except Exception as e:
                explanation = f"Error generating description: {e}"

            logic_steps.append({
                "Step": step_number,
                "Named Range": name,
                "Description": explanation
            })

        # If no _cN_ logic blocks found, issue warning
        if not logic_steps:
            st.warning("⚠️ No logic components found using `_c1_`, `_c2_`, etc. naming convention. Please check that named ranges follow this format.")

        # --- Checks and Validation ---
        check_pattern = re.compile(r"^_ch(\d+)_.*")  # matches _ch1_, _ch2_, etc.
        check_summaries = {}

        # Extract and order checks
        for name in summaries:
            match = check_pattern.match(name)
            if match:
                check_num = int(match.group(1))
                check_summaries[check_num] = name

        check_data = []
        for check_num in sorted(check_summaries):
            name = check_summaries[check_num]
            summary_json = summaries[name]
            json_summary = summary_json.get("summary", "")
            general_formula = summary_json.get("general_formula", "")
            sheet = summary_json.get("sheet_name", "")
            excel_range = summary_json.get("excel_range", "")

            # GPT prompt to describe what the check does
            prompt = f"""You are an expert actuary and spreadsheet modeller.

        You are reviewing a check step in an Excel model using the Lee-Carter mortality framework.

        The named range is `{name}`, located in `{sheet}` `{excel_range}`.

        Here is a description of the logic:
        "{json_summary}"

        And the general formula pattern:
        "{general_formula}"

        What is this check checking for? Summarize the logic of the check in 1–2 confident sentences and state whether it appears to validate a model result, flag an inconsistency, or confirm an assumption.

        Avoid vague language. Be concise and clear."""

            try:
                response = client.chat.completions.create(
                    model="gpt-4",
                    messages=[
                        {"role": "system", "content": "You describe spreadsheet checks in actuarial models."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.3
                )
                description = response.choices[0].message.content.strip()
            except Exception as e:
                description = f"Error: {e}"

            check_data.append({
                "Check No.": check_num,
                "Named Range": name,
                "Description": description
            })

        # Convert to DataFrame for Streamlit and Word doc
        checks_df = pd.DataFrame(check_data)

        # --- Assumptions and Limitations ---
        try:
            all_summaries = "\n".join(
                f"{k}: {v.get('summary', '')}" for k, v in summaries.items() if "summary" in v
            )
            all_formulas = "\n".join(
                f"{k}: {v.get('general_formula', '')}" for k, v in summaries.items() if "general_formula" in v
            )

            assumptions_prompt = f"""You are an expert actuary and spreadsheet modeller reviewing a workbook based on the Lee-Carter mortality model.

        Below are summaries and formulas used in various named ranges of the model:

        Summaries:
        {all_summaries}

        Formulas:
        {all_formulas}

        From this information, write a concise paragraph that lists the **key assumptions** underlying this spreadsheet model (e.g., trends, input behaviour, mortality evolution), and any **notable limitations or simplifications** (e.g., no sensitivity testing, static inputs, deterministic projections).

        Avoid vague language like "possibly" or "might". Use confident and professional actuarial language."""

            assumptions_response = client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You describe assumptions and limitations in actuarial spreadsheet models."},
                    {"role": "user", "content": assumptions_prompt}
                ],
                temperature=0.3
            )
            assumptions_text = assumptions_response.choices[0].message.content.strip()

        except Exception as e:
            assumptions_text = f"Error generating assumptions and limitations: {e}"
        
        with st.expander("📄 Spreadsheet Document", expanded=False):
            st.title("📄 Model Documentation")

            st.header("## Version Control")

            # Model version control table
            st.subheader("### Model Version Control")
            model_version_df = pd.DataFrame({
                "Version": [""],
                "Date": ["00/00/0000"],
                "Info": [""],
                "Updated by": [""],
                "Reviewed by": [""],
                "Review Date": [""]
            })
            st.dataframe(model_version_df, use_container_width=True)

            # Documentation version control table
            st.subheader("### Documentation Version Control")
            doc_version_df = pd.DataFrame({
                "Version": [""],
                "Date": ["00/00/0000"],
                "Info": [""],
                "Updated by": [""],
                "Reviewed by": [""],
                "Review Date": [""]
            })
            st.dataframe(doc_version_df, use_container_width=True)

            # Ownership section
            st.header("## Ownership")
            st.text("### Owner")
            st.text("### Risk rating (or other client control standard)")
            st.text("### Internal audit history")

            # Purpose
            st.header("## Purpose")
            st.text_area("Describe the purpose of the model:", value=model_purpose)

            # Inputs table
            st.header("## Inputs")
            row_height = 35
            max_height = 500
            calculated_height = min(len(inputs_df) * row_height + 35, max_height)
            st.dataframe(inputs_df, use_container_width=True)

            # Outputs
            st.header("## Outputs")
            st.dataframe(outputs_df, use_container_width=True)

            # Logic
            st.header("## Logic")
            if logic_steps:
                logic_df = pd.DataFrame(logic_steps)
                st.dataframe(logic_df, use_container_width=True)
            else:
                st.warning("⚠️ No logic documented. Ensure named ranges follow `_c1_`, `_c2_`, ... convention.")

            # Checks and validation
            st.header("## Checks and Validation")
            if not checks_df.empty:
                st.dataframe(checks_df, use_container_width=True)
            else:
                st.warning("⚠️ No checks found using `_ch1_`, `_ch2_`, etc. naming convention. The model may not include validation steps.")

            # Assumptions and limitations
            st.header("## Assumptions and Limitations")
            st.text_area("List assumptions and limitations:", value=assumptions_text, height=200)

            # TAS Compliance
            st.header("## TAS Compliance")
            st.text_area("Describe how the model complies with TAS:")

        # JSON download
        json_str = json.dumps(summaries, indent=2)
        st.download_button("📥 Download JSON Summary", data=json_str, file_name="named_range_summaries.json", mime="application/json")


        doc = Document()
        doc.add_heading("Named Range JSON Summary", 0)
        for name, summary in summaries.items():
            doc.add_heading(name, level=1)
            for key, value in summary.items():
                if isinstance(value, (list, dict)):
                    value = json.dumps(value, indent=2)
                doc.add_paragraph(f"{key}: {value}")

        # 🧾 Add Spreadsheet Document content to Word
        doc.add_page_break()
        doc.add_heading("📄 Spreadsheet Documentation", 0)

        # Version Control
        doc.add_heading("Version Control", level=1)
        doc.add_heading("Model Version Control", level=2)
        for col in model_version_df.columns:
            doc.add_paragraph(f"{col}: __________")

        doc.add_heading("Documentation Version Control", level=2)
        for col in doc_version_df.columns:
            doc.add_paragraph(f"{col}: __________")

        # Ownership
        doc.add_heading("Ownership", level=1)
        doc.add_paragraph("Owner: __________")
        doc.add_paragraph("Risk rating (or other client control standard): __________")
        doc.add_paragraph("Internal audit history: __________")

        # Purpose
        doc.add_heading("Purpose", level=1)
        doc.add_paragraph(model_purpose)

        # Inputs Table
        # Inputs Table (formatted as a Word table)
        doc.add_heading("Inputs", level=1)
        table = doc.add_table(rows=1, cols=5)
        table.autofit = True
        table.style = "Table Grid"

        # Add table headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "No."
        hdr_cells[1].text = "Name"
        hdr_cells[2].text = "Type"
        hdr_cells[3].text = "Source"
        hdr_cells[4].text = "Info"

        # Add data rows
        for row in inputs_data:
            row_cells = table.add_row().cells
            row_cells[0].text = str(row["No."])
            row_cells[1].text = row["Name"]
            row_cells[2].text = row["Type"]
            row_cells[3].text = row["Source"]
            row_cells[4].text = row["Info"]

        # Other sections
        doc.add_heading("Outputs", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.autofit = True
        table.style = "Table Grid"

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "No."
        hdr_cells[1].text = "Name"
        hdr_cells[2].text = "Description"

        for row in outputs_data:
            row_cells = table.add_row().cells
            row_cells[0].text = str(row["No."])
            row_cells[1].text = row["Name"]
            row_cells[2].text = row["Description"]

        doc.add_heading("Logic", level=1)
        if logic_steps:
            table = doc.add_table(rows=1, cols=3)
            table.autofit = True
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Step"
            hdr_cells[1].text = "Named Range"
            hdr_cells[2].text = "Description"

            for row in logic_steps:
                row_cells = table.add_row().cells
                row_cells[0].text = str(row["Step"])
                row_cells[1].text = row["Named Range"]
                row_cells[2].text = row["Description"]
        else:
            doc.add_paragraph("⚠ No logic components found with the expected `_cN_` naming pattern.")

        doc.add_heading("Checks and Validation", level=1)
        if not checks_df.empty:
            table = doc.add_table(rows=1, cols=3)
            table.autofit = True
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Check No."
            hdr_cells[1].text = "Named Range"
            hdr_cells[2].text = "Description"

            for row in check_data:
                row_cells = table.add_row().cells
                row_cells[0].text = str(row["Check No."])
                row_cells[1].text = row["Named Range"]
                row_cells[2].text = row["Description"]
        else:
            doc.add_paragraph("⚠ No validation checks found using `_chN_` naming pattern.")

        doc.add_heading("Assumptions and Limitations", level=1)
        doc.add_paragraph(assumptions_text)

        doc.add_heading("TAS Compliance", level=1)
        doc.add_paragraph("Describe how the model complies with TAS:")

        
        docx_io = BytesIO()
        doc.save(docx_io)
        docx_io.seek(0)

        st.download_button(
            "📄 Download Summary as Word Document",
            data=docx_io,
            file_name="named_range_summary.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    else:
        st.info("Press the button above to generate a GPT-based JSON and Documentation.")

else:
    st.info("⬆️ Upload one or more .xlsx files to begin.")
